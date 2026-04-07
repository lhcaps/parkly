"""
Parkly v12 DOCX Snapshot Generator
Generates professional Word documents from markdown source files.
"""

import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Palette ─────────────────────────────────────────────────────────────────
DARK_NAVY  = RGBColor(0x0F, 0x1E, 0x3D)   # headings, borders
ACCENT     = RGBColor(0x1D, 0x4E, 0xD6)   # accent bars, links
MID_GRAY   = RGBColor(0x55, 0x65, 0x7A)   # sub-headings
LIGHT_BG   = RGBColor(0xF1, 0xF5, 0xF9)   # code block background
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
RULE_COLOR = RGBColor(0x1D, 0x4E, 0xD6)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_bottom_border(paragraph, color_hex="1D4ED6", size=12):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_spacing(para, before=0, after=6):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)

def add_heading_bar(doc, text, level=1):
    """Accent-bar heading with bold label."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=14, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "1D4ED6")
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_NAVY
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_NAVY
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    elif level == 3:
        run.font.size = Pt(11)
    else:
        run.font.size = Pt(10)
    return p

def add_subheading(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=2)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(10)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, before=1, after=2)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def add_code_block(doc, code_text):
    lines = code_text.strip().split("\n")
    for line in lines:
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)
        p.paragraph_format.left_indent = Inches(0.25)
        # Light gray shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F1F5F9")
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    doc.add_paragraph()  # spacer

def add_divider(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    add_bottom_border(p, "1D4ED6", 6)

def add_table(doc, headers, rows, col_widths=None):
    """Styled table with navy header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "0F1E3D")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else "F8FAFC"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Column widths
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    doc.add_paragraph()
    return table

def add_json_block(doc, json_text):
    """Coloured JSON code block."""
    lines = json_text.strip().split("\n")
    for line in lines:
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)
        p.paragraph_format.left_indent = Inches(0.25)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E293B")
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    doc.add_paragraph()

def add_info_box(doc, label, text, fill="EBF4FF"):
    """Coloured info/callout box."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    if label:
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        r1.font.color.rgb = ACCENT
        r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

# ─── Markdown Parser ──────────────────────────────────────────────────────────

def parse_markdown(doc, md_text: str):
    """Convert markdown string to formatted docx paragraphs."""
    lines = md_text.split("\n")
    i = 0
    in_code = False
    in_json = False
    in_table = False
    table_buf = []
    pending_h2 = None

    def flush_table():
        nonlocal table_buf
        if table_buf:
            headers = table_buf[0]
            rows = table_buf[1:]
            if headers and rows:
                add_table(doc, headers, rows)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                # Check for json
                lang = line.strip()[3:].strip()
                if lang in ("json", "jsonc"):
                    in_json = True
                i += 1
                continue
            else:
                in_code = False
                in_json = False
                i += 1
                continue

        if in_code:
            if in_json:
                add_json_block(doc, line)
            else:
                add_code_block(doc, line)
            i += 1
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            stripped = line.strip()
            # Divider row
            if set(stripped.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set():
                i += 1
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            # Check if first cell looks like a header
            if i == 0 or (i > 0 and "|" in lines[i-1] and lines[i-1].strip().startswith("|") and set(lines[i-1].replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) != set()):
                table_buf.append(cells)
            elif table_buf:
                table_buf.append(cells)
            else:
                table_buf.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        # H1
        if line.startswith("# "):
            add_heading_bar(doc, line[2:].strip(), level=1)
        # H2
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
        # H3
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
        # H4
        elif line.startswith("#### "):
            add_subheading(doc, line[5:].strip())
        # HR
        elif line.strip() in ("---", "***", "___"):
            add_divider(doc)
        # Bullet
        elif line.strip().startswith("- "):
            text = line.strip()[2:]
            # sub-bullet
            if text.startswith("  "):
                add_bullet(doc, text.strip().lstrip("- ").lstrip("* "), level=1)
            else:
                add_bullet(doc, text)
        elif line.strip().startswith("* "):
            add_bullet(doc, line.strip()[2:])
        # Empty
        elif line.strip() == "":
            p = doc.add_paragraph()
            set_para_spacing(p, before=0, after=2)
        # Body / metadata
        else:
            text = line.strip()
            if text.startswith("|") or text.startswith("...") or text.startswith("```"):
                pass
            elif re.match(r"^[-*+] ", text):
                pass
            else:
                # Inline bold/italic
                add_body(doc, text)

        i += 1

    if in_table:
        flush_table()

# ─── Document Factories ───────────────────────────────────────────────────────

def make_cover(doc, title, subtitle, date_str, baseline=""):
    # Top accent bar
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "0F1E3D")
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(" " * 80)
    run.font.size = Pt(4)

    doc.add_paragraph()

    # Title
    tp = doc.add_paragraph()
    set_para_spacing(tp, before=24, after=8)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = DARK_NAVY

    # Subtitle
    sp = doc.add_paragraph()
    set_para_spacing(sp, before=4, after=4)
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.font.size = Pt(14)
    sr.font.color.rgb = ACCENT

    doc.add_paragraph()

    # Meta table
    meta_rows = [
        ("Snapshot Date", date_str),
        ("Release Baseline", baseline),
        ("Version", "v12"),
        ("Status", "Approved"),
    ]
    add_table(doc, ["Field", "Value"],
              [[k, v] for k, v in meta_rows],
              col_widths=[1.8, 4.2])

    # Bottom bar
    p2 = doc.add_paragraph()
    pPr2 = p2._p.get_or_add_pPr()
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:val"), "clear")
    shd2.set(qn("w:color"), "auto")
    shd2.set(qn("w:fill"), "1D4ED6")
    pPr2.append(shd2)
    p2.paragraph_format.space_before = Pt(18)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run(" " * 80)
    r2.font.size = Pt(6)

    doc.add_page_break()

def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)

def make_snapshot_id_block(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EFF6FF")
    pPr.append(shd)
    r1 = p.add_run("SNAPSHOT  ")
    r1.bold = True
    r1.font.color.rgb = ACCENT
    r1.font.size = Pt(9)
    r2 = p.add_run("This document is a point-in-time snapshot of the Parkly repository. It captures the current shipped architecture, security posture, data model, deployment shape, observability, and quality gate.")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ─── File Generators ──────────────────────────────────────────────────────────

def gen_project_snapshot():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_Project_Snapshot_v12_{timestamp}.docx"
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "Enterprise B2B SaaS Parking Management System\nProject Technical Snapshot v12",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    make_snapshot_id_block(doc)
    add_divider(doc)

    md = open("docs/SPEC-v11.md", encoding="utf-8").read()
    parse_markdown(doc, md)

    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

def gen_runbook():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_Runbook_v12_{timestamp}.docx"
    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "Operations Runbook\nRelease Procedures, Incident Handling & Recovery",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    add_divider(doc)
    md = open("docs/RUNBOOK.md", encoding="utf-8").read()
    parse_markdown(doc, md)
    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

def gen_api_doc():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_API_Contract_v12_{timestamp}.docx"
    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "API Contract Specification\nAuth, Gate, Dashboard, Topology, Subscriptions, and Realtime Surfaces",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    make_snapshot_id_block(doc)
    add_divider(doc)
    md = open("docs/API.md", encoding="utf-8").read()
    parse_markdown(doc, md)
    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

def gen_architecture():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_Architecture_v12_{timestamp}.docx"
    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "Architecture Document\nRuntime Topology, Bounded Contexts, and Deployment Profiles",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    make_snapshot_id_block(doc)
    add_divider(doc)
    md = open("docs/ARCHITECTURE.md", encoding="utf-8").read()
    parse_markdown(doc, md)
    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

def gen_evidence():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_Release_Evidence_v12_{timestamp}.docx"
    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "Release Evidence Standard\nCI, QA, Deployment Readiness & Sign-Off Criteria",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    make_snapshot_id_block(doc)
    add_divider(doc)
    md = open("docs/EVIDENCE.md", encoding="utf-8").read()
    parse_markdown(doc, md)
    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

def gen_adr():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_ADR_v12_{timestamp}.docx"
    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "Architecture Decision Records\nCore Platform Decisions & Rationale",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    add_divider(doc)
    md = open("docs/ADR.md", encoding="utf-8").read()
    parse_markdown(doc, md)
    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

def gen_error_codes():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_Error_Catalog_v12_{timestamp}.docx"
    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "Error Catalog\nCanonical API Error Codes and Response Shapes",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    add_divider(doc)
    md = open("docs/ERROR_CODES.md", encoding="utf-8").read()
    parse_markdown(doc, md)
    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

def gen_rbac():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_RBAC_Matrix_v12_{timestamp}.docx"
    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "RBAC Matrix\nCanonical Roles, API Access Matrix, and Web Route Policy",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    add_divider(doc)
    md = open("docs/RBAC_MATRIX_API.md", encoding="utf-8").read()
    parse_markdown(doc, md)
    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

def gen_retention_policy():
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path  = f"docs/spec/Parkly_Retention_Policy_v12_{timestamp}.docx"
    doc = Document()
    set_doc_defaults(doc)
    make_cover(doc,
               "PARKLY",
               "Retention Policy\nData Cleanup, Demo Preservation, and Audit Protection",
               datetime.now().strftime("%d %B %Y"),
               "backend-rc1")
    add_divider(doc)
    md = open("docs/RETENTION_POLICY.md", encoding="utf-8").read()
    parse_markdown(doc, md)
    doc.save(out_path)
    print(f"[OK] {out_path}")
    return out_path

if __name__ == "__main__":
    paths = []
    paths.append(gen_project_snapshot())
    paths.append(gen_runbook())
    paths.append(gen_api_doc())
    paths.append(gen_architecture())
    paths.append(gen_evidence())
    paths.append(gen_adr())
    paths.append(gen_error_codes())
    paths.append(gen_rbac())
    paths.append(gen_retention_policy())

    print("\nAll documents generated:")
    for p in paths:
        print(f"  {p}")
