"""Validate docx structure and content."""
import glob
from docx import Document

files = sorted(glob.glob("docs/spec/Parkly_*_v12_*.docx"))
all_ok = True
for fpath in files:
    doc = Document(fpath)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = doc.tables
    total_words = sum(len(p.split()) for p in paras)
    print(f"OK  {fpath.split('/')[-1]}")
    print(f"    paragraphs={len(paras)}, tables={len(tables)}, words={total_words}")
    if total_words < 100:
        print(f"    WARN: suspiciously few words!")
        all_ok = False
    # Check cover title
    if not any("PARKLY" in p for p in paras[:5]):
        print(f"    WARN: no PARKLY title found in first 5 paragraphs")
        all_ok = False

print()
if all_ok:
    print("All documents validated successfully.")
else:
    print("Some documents have warnings.")